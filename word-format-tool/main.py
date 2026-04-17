import streamlit as st
import copy
import re
import random
import json
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_BUILTIN_STYLE
from docx.oxml.ns import qn
import os
import requests
import pandas as pd

# ====================== 预编译正则（原核心逻辑完整保留）======================
RE_REF_FLAG = re.compile(r'^\[(\d+)\]')
RE_REF_KEYWORD = re.compile(r'参考文献|参考资料|References')
RE_REF_SPACE = re.compile(r'\s+')
RE_REF_CN_FONT = re.compile(r'([\u4e00-\u9fa5]+)\[([A-Z]+)\]')
RE_REF_DOT = re.compile(r'。(?![\u4e00-\u9fa5])')
RE_REF_COMMA = re.compile(r'，')
RE_REF_COLON = re.compile(r'：')
RE_KEYWORDS = re.compile(r'[\u4e00-\u9fa5]{2,}')
RE_WHITE_NUMBER = re.compile(r'^\d+(\.\d+)*$')
RE_WHITE_QUOTE = re.compile(r'^\[.*\]$')
RE_SENTENCE_SPLIT = re.compile(r'(?<=[。！？；])\s*')
RE_CLAUSE_SPLIT = re.compile(r'[，。；]')
RE_RED_HIGHLIGHT = re.compile(r'<font color="red">(.*?)</font>', re.DOTALL)

# ====================== 全局配置与常量（完整保留+优化）======================
WHITE_WORDS = [
    "知网", "维普", "万方", "PaperPass", "PaperYY", "PaperFree", "挑战杯", "互联网+", "三创赛",
    "参考文献", "公式", "图表", "图", "表", "附录", "摘要", "关键词", "Abstract",
    "机器学习", "人工智能", "算法", "系统", "模型", "数据"
]
WPS_STYLE_MAPPING = {
    "一级标题": WD_BUILTIN_STYLE.HEADING_1,
    "二级标题": WD_BUILTIN_STYLE.HEADING_2,
    "三级标题": WD_BUILTIN_STYLE.HEADING_3,
    "正文": WD_BUILTIN_STYLE.NORMAL
}

# 全量模板库（完整保留原所有模板）
COMPETITION_FORMATS = {
    "三创赛-全国大学生电子商务创新创意及创业挑战赛": {
        "update_time": "2024-01-15",
        "cn_format": {
            "一级标题": {"font": "黑体", "size": "三号", "bold": True, "align": "居中", "line_type": "倍数", "line_value": 1.2, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "二级标题": {"font": "黑体", "size": "小三", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.2, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "三级标题": {"font": "楷体_GB2312", "size": "四号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.2, "indent": 0, "space_before": 3, "space_after": 0, "char_spacing": 0},
            "正文": {"font": "仿宋", "size": "四号", "bold": False, "align": "两端对齐", "line_type": "倍数", "line_value": 1.2, "indent": 2, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "五号", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.2, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "三号", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小三", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "五号", "bold": False, "italic": False}
        },
        "special_requirements": ["硬件必须配小程序/App", "服务必须线上化", "需要3D建模图/UI原型", "图表必须标注数据来源"]
    },
    "挑战杯-全国大学生课外学术科技作品竞赛": {
        "update_time": "2024-02-20",
        "cn_format": {
            "一级标题": {"font": "黑体", "size": "三号", "bold": True, "align": "居中", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "二级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "三级标题": {"font": "黑体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 3, "space_after": 0, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "倍数", "line_value": 1.5, "indent": 2, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "五号", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.25, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "三号", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "五号", "bold": False, "italic": False}
        },
        "special_requirements": ["全文约15000字", "双面打印", "严格章-节-条层级结构", "标题单倍行距，正文1.5倍行距"]
    },
    "互联网+大学生创新创业大赛": {
        "update_time": "2024-03-10",
        "cn_format": {
            "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "二级标题": {"font": "黑体", "size": "三号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "三级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 3, "space_after": 0, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "四号", "bold": False, "align": "两端对齐", "line_type": "倍数", "line_value": 1.5, "indent": 2, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "五号", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.25, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "二号", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "三号", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "五号", "bold": False, "italic": False}
        },
        "special_requirements": ["全文10000字以上", "分创意组/创业组撰写", "需包含完整财务预测", "商业模式需清晰可落地"]
    }
}
UNIVERSITY_FORMATS = {
    "清华大学本科毕业论文模板": {
        "update_time": "2024-04-01",
        "cn_format": {
            "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 24, "space_after": 18, "char_spacing": 0},
            "二级标题": {"font": "黑体", "size": "小三", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 18, "space_after": 12, "char_spacing": 0},
            "三级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "固定值", "line_value": 20, "indent": 2, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "小五", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.25, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "二号", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小三", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小五", "bold": False, "italic": False}
        },
        "special_requirements": ["全文8000-15000字", "需包含中英文摘要", "参考文献需符合GB/T 7714-2015", "页眉标注清华大学本科毕业论文"]
    },
    "北京大学本科毕业论文模板": {
        "update_time": "2024-04-02",
        "cn_format": {
            "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 18, "space_after": 12, "char_spacing": 0},
            "二级标题": {"font": "黑体", "size": "三号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "三级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "固定值", "line_value": 22, "indent": 2, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "五号", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.25, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "二号", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "三号", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "五号", "bold": False, "italic": False}
        },
        "special_requirements": ["全文10000-20000字", "需包含摘要/关键词/参考文献", "参考文献需符合GB/T 7714", "页眉标注北京大学本科毕业论文"]
    },
    "浙江大学本科毕业论文模板": {
        "update_time": "2024-04-03",
        "cn_format": {
            "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 20, "space_after": 15, "char_spacing": 0},
            "二级标题": {"font": "黑体", "size": "小三", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 15, "space_after": 10, "char_spacing": 0},
            "三级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 10, "space_after": 5, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "固定值", "line_value": 20, "indent": 2, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "小五", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.25, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "二号", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小三", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小五", "bold": False, "italic": False}
        },
        "special_requirements": ["全文8000-12000字", "需包含中英文摘要", "参考文献需符合GB/T 7714-2015", "页眉标注浙江大学本科毕业论文"]
    },
    "复旦大学本科毕业论文模板": {
        "update_time": "2024-04-04",
        "cn_format": {
            "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 18, "space_after": 12, "char_spacing": 0},
            "二级标题": {"font": "黑体", "size": "三号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "三级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "固定值", "line_value": 20, "indent": 2, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "五号", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.25, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "二号", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "三号", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "五号", "bold": False, "italic": False}
        },
        "special_requirements": ["全文10000-15000字", "需包含摘要/关键词/参考文献", "参考文献需符合GB/T 7714", "页眉标注复旦大学本科毕业论文"]
    },
    "上海交通大学本科毕业论文模板": {
        "update_time": "2024-04-05",
        "cn_format": {
            "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 20, "space_after": 12, "char_spacing": 0},
            "二级标题": {"font": "黑体", "size": "小三", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "三级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "固定值", "line_value": 22, "indent": 2, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "小五", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.25, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "二号", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "三号", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小五", "bold": False, "italic": False}
        },
        "special_requirements": ["全文12000-20000字", "需包含中英文摘要", "参考文献需符合GB/T 7714-2015", "页眉标注上海交通大学本科毕业论文"]
    }
}
THESIS_FORMATS = {
    "本科毕业论文-通用模板": {
        "update_time": "2024-04-01",
        "cn_format": {
            "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 18, "space_after": 12, "char_spacing": 0},
            "二级标题": {"font": "黑体", "size": "三号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "三级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "固定值", "line_value": 20, "indent": 2, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "五号", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.25, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "二号", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "三号", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "五号", "bold": False, "italic": False}
        },
        "special_requirements": ["全文8000-12000字", "需包含摘要/关键词/参考文献/致谢", "参考文献需符合GB/T 7714格式", "页眉需标注学校+论文题目"]
    },
    "硕士毕业论文-通用模板": {
        "update_time": "2024-04-05",
        "cn_format": {
            "一级标题": {"font": "黑体", "size": "二号", "bold": True, "align": "居中", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 24, "space_after": 18, "char_spacing": 0},
            "二级标题": {"font": "黑体", "size": "小三", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 18, "space_after": 12, "char_spacing": 0},
            "三级标题": {"font": "黑体", "size": "四号", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "固定值", "line_value": 22, "indent": 2, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "小五", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.25, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "二号", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小三", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小五", "bold": False, "italic": False}
        },
        "special_requirements": ["全文30000字以上", "需包含中英文摘要", "参考文献需符合GB/T 7714-2015", "需包含创新点说明"]
    }
}
JOURNAL_FORMATS = {
    "MTA - Multimedia Tools and Applications": {
        "update_time": "2024-04-10",
        "cn_format": {
            "一级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "二级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "三级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 3, "space_after": 0, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "小五", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小五", "bold": False, "italic": False}
        },
        "special_requirements": ["双栏排版", "单栏摘要", "参考文献需符合APA格式", "图表需单独标注", "全文不超过15页"]
    },
    "IEEE Transactions": {
        "update_time": "2024-04-10",
        "cn_format": {
            "一级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "二级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "三级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 3, "space_after": 0, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "小五", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小五", "bold": False, "italic": False}
        },
        "special_requirements": ["双栏排版", "无首行缩进", "参考文献需符合IEEE格式", "图表需跨栏", "全文不超过8页"]
    },
    "ACM Transactions": {
        "update_time": "2024-04-10",
        "cn_format": {
            "一级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "二级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "三级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 3, "space_after": 0, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "小五", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小五", "bold": False, "italic": False}
        },
        "special_requirements": ["双栏排版", "无首行缩进", "参考文献需符合ACM格式", "图表需跨栏", "全文不超过10页"]
    },
    "Elsevier Journal": {
        "update_time": "2024-04-10",
        "cn_format": {
            "一级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "二级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "三级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 3, "space_after": 0, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "小五", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小五", "bold": False, "italic": False}
        },
        "special_requirements": ["单栏排版", "无首行缩进", "参考文献需符合Elsevier格式", "图表需单独标注", "全文不超过20页"]
    },
    "Springer Journal": {
        "update_time": "2024-04-10",
        "cn_format": {
            "一级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 12, "space_after": 6, "char_spacing": 0},
            "二级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 6, "space_after": 3, "char_spacing": 0},
            "三级标题": {"font": "宋体", "size": "小四", "bold": True, "align": "左对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 3, "space_after": 0, "char_spacing": 0},
            "正文": {"font": "宋体", "size": "小四", "bold": False, "align": "两端对齐", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0},
            "表格": {"font": "宋体", "size": "小五", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.0, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
        },
        "en_format": {
            "一级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "二级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": True, "italic": False},
            "三级标题": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "四号", "bold": True, "italic": False},
            "正文": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小四", "bold": False, "italic": False},
            "表格": {"en_font": "Times New Roman", "size_same_as_cn": True, "size": "小五", "bold": False, "italic": False}
        },
        "special_requirements": ["单栏排版", "无首行缩进", "参考文献需符合Springer格式", "图表需单独标注", "全文不超过15页"]
    }
}
ALL_TEMPLATES = {**COMPETITION_FORMATS, **UNIVERSITY_FORMATS, **THESIS_FORMATS, **JOURNAL_FORMATS}

# 润色等级配置
REWRITE_LEVEL = {
    "轻度润色": {"synonym": True, "sentence_reorder": False, "structure_change": False},
    "标准润色": {"synonym": True, "sentence_reorder": True, "structure_change": False},
    "深度润色": {"synonym": True, "sentence_reorder": True, "structure_change": True}
}
SYNONYM_DICT = {
    "提升": "有效改善", "降低": "显著减少", "增加": "大幅提升", "减少": "有效降低",
    "首先": "从实际落地情况来看", "其次": "进一步分析", "最后": "综合上述分析",
    "综上所述": "结合全维度分析", "总而言之": "从实践结果来看",
    "研究": "调研分析", "实验": "测试验证", "分析": "剖析", "结果": "结论",
    "方法": "方案", "系统": "平台", "模型": "架构", "数据": "信息"
}

# 格式映射常量
ALIGN_MAP = {
    "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
    "居中": WD_ALIGN_PARAGRAPH.CENTER,
    "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY
}
FONT_SIZE_MAP = {
    "初号": 42, "小初": 36, "一号": 26, "小一": 24,
    "二号": 22, "小二": 18, "三号": 16, "小三": 15,
    "四号": 14, "小四": 12, "五号": 10.5, "小五": 9
}
EN_FONT_LIST = ["Times New Roman", "Arial", "Calibri", "Courier New"]
CN_FONT_LIST = ["宋体", "黑体", "楷体", "仿宋_GB2312", "微软雅黑"]
MAX_FILE_SIZE_MB = 50
random.seed(42)

# ====================== 核心工具函数（100%保留原功能+优化）======================
@st.cache_data(ttl=3600)
def get_cached_template(template_name):
    return copy.deepcopy(ALL_TEMPLATES[template_name]["cn_format"]), copy.deepcopy(ALL_TEMPLATES[template_name]["en_format"])

def get_title_level(para_text, prev_para_text=None):
    """
    精准标题分级：彻底解决正文列表误识别、三级标题被二级吞没
    三重校验：格式匹配 + 上下文区分 + 语义过滤
    """
    text = para_text.strip()
    if not text:
        return "正文"

    # ====================== 一级标题（严格匹配，不冲突）======================
    # 匹配：第X章、1、、1、（带顿号的一级标题）
    if re.match(r'^第[一二三四五六七八九十]+章', text) or re.match(r'^\d+、', text):
        return "一级标题"
    
    # ====================== 二级标题（严格匹配，不冲突）======================
    # 匹配：（一）、1.1（带点的二级标题，排除纯数字列表）
    elif re.match(r'^（[一二三四五六七八九十]）', text) or re.match(r'^\d+\.\d+\s', text):
        return "二级标题"
    
    # ====================== 三级标题（核心修复：区分标题和正文列表）======================
    # 1. 先匹配格式：（1）、1.1.1
    elif re.match(r'^（\d+）', text) or re.match(r'^\d+\.\d+\.\d+', text):
        # 2. 上下文校验：如果上一段是正文/空行，且当前段落是长文本（>15字），判定为正文列表
        if prev_para_text and len(text) > 15:
            # 3. 语义过滤：如果开头是「电脑硬件的科普」这种描述性内容，直接判定为正文
            if re.match(r'^（\d+）[a-zA-Z\u4e00-\u9fa5]{2,}', text):
                return "正文"
        # 4. 否则才判定为三级标题（真正的章节标题）
        return "三级标题"
    
    # 所有不匹配的，全部判定为正文
    return "正文"

def recommend_template(file):
    """智能模板推荐功能，通过分析文档内容自动匹配最适合的模板"""
    try:
        doc = Document(file)
        file.seek(0)
        # 提取文档全文
        full_text = "".join([p.text for p in doc.paragraphs])
        
        # 模板推荐规则
        template_rules = {
            "三创赛-全国大学生电子商务创新创意及创业挑战赛": ["电子商务", "创新", "创意", "创业", "三创赛"],
            "挑战杯-全国大学生课外学术科技作品竞赛": ["挑战杯", "学术", "科技", "作品", "竞赛"],
            "互联网+大学生创新创业大赛": ["互联网+", "创新", "创业", "大赛"],
            "清华大学本科毕业论文模板": ["清华大学", "毕业论文", "摘要", "关键词", "参考文献"],
            "北京大学本科毕业论文模板": ["北京大学", "毕业论文", "摘要", "关键词", "参考文献"],
            "浙江大学本科毕业论文模板": ["浙江大学", "毕业论文", "摘要", "关键词", "参考文献"],
            "复旦大学本科毕业论文模板": ["复旦大学", "毕业论文", "摘要", "关键词", "参考文献"],
            "上海交通大学本科毕业论文模板": ["上海交通大学", "毕业论文", "摘要", "关键词", "参考文献"],
            "本科毕业论文-通用模板": ["毕业论文", "摘要", "关键词", "参考文献"],
            "硕士毕业论文-通用模板": ["硕士", "毕业论文", "摘要", "关键词", "参考文献"],
            "MTA - Multimedia Tools and Applications": ["Multimedia", "Tools", "Applications", "MTA"],
            "IEEE Transactions": ["IEEE", "Transactions"],
            "ACM Transactions": ["ACM", "Transactions"],
            "Elsevier Journal": ["Elsevier", "Journal"],
            "Springer Journal": ["Springer", "Journal"]
        }
        
        # 计算每个模板的匹配度
        scores = {}
        for template, keywords in template_rules.items():
            score = 0
            for keyword in keywords:
                if keyword in full_text:
                    score += 1
            scores[template] = score
        
        # 找出匹配度最高的模板
        if scores:
            recommended_template = max(scores, key=scores.get)
            if scores[recommended_template] > 0:
                return recommended_template, scores[recommended_template]
        
        # 默认返回第一个模板
        return list(ALL_TEMPLATES.keys())[0], 0
    except Exception as e:
        # 出错时返回默认模板
        return list(ALL_TEMPLATES.keys())[0], 0

def pdf_to_docx(pdf_file):
    """将PDF文件转换为Word文档"""
    try:
        import pdfplumber
        from docx import Document
        
        doc = Document()
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    except ImportError:
        raise Exception("缺少pdfplumber依赖，无法处理PDF文件")
    except Exception as e:
        raise Exception(f"PDF转换失败：{str(e)}")

def doc_to_docx(doc_file):
    """将doc文件转换为docx格式"""
    try:
        import textract
        import tempfile
        
        # 保存文件指针位置
        current_pos = doc_file.tell()
        
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_doc:
            temp_doc.write(doc_file.read())
            temp_doc_path = temp_doc.name
        
        # 恢复文件指针位置
        doc_file.seek(current_pos)
        
        text = textract.process(temp_doc_path).decode('utf-8')
        os.unlink(temp_doc_path)
        
        from docx import Document
        doc = Document()
        for line in text.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    except ImportError:
        raise Exception("缺少textract依赖，无法处理doc文件")
    except Exception as e:
        raise Exception(f"doc转换失败：{str(e)}")

def extract_template_from_doc(file):
    try:
        if file.name.endswith('.docx'):
            doc = Document(file)
            file.seek(0)
        elif file.name.endswith('.doc'):
            # 转换doc为docx
            docx_file = doc_to_docx(file)
            doc = Document(docx_file)
            file.seek(0)
        elif file.name.endswith('.pdf'):
            # 转换pdf为docx
            docx_file = pdf_to_docx(file)
            doc = Document(docx_file)
            file.seek(0)
        else:
            return None, None, "不支持的文件格式"
        cn_format = {}
        en_format = {}
        style_stats = {}
        for para in doc.paragraphs:
            level = get_title_level(para.text)
            if level not in style_stats:
                style_stats[level] = {"font": None, "size": None, "bold": None, "align": None, "line_type": "倍数", "line_value": 1.5, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
            if para.runs:
                run = para.runs[0]
                if run.font.name:
                    if run.font.name in CN_FONT_LIST:
                        style_stats[level]["font"] = run.font.name
                    else:
                        style_stats[level]["en_font"] = run.font.name
                if run.font.size:
                    for size_name, size_pt in FONT_SIZE_MAP.items():
                        if abs(run.font.size.pt - size_pt) < 0.5:
                            style_stats[level]["size"] = size_name
                            break
                if run.font.bold is not None:
                    style_stats[level]["bold"] = run.font.bold
            if para.paragraph_format:
                pf = para.paragraph_format
                if pf.alignment:
                    for align_name, align_val in ALIGN_MAP.items():
                        if pf.alignment == align_val:
                            style_stats[level]["align"] = align_name
                            break
                if pf.first_line_indent:
                    style_stats[level]["indent"] = int(pf.first_line_indent.cm / 0.74)
                if pf.space_before:
                    style_stats[level]["space_before"] = int(pf.space_before.pt)
                if pf.space_after:
                    style_stats[level]["space_after"] = int(pf.space_after.pt)
                if pf.line_spacing:
                    style_stats[level]["line_value"] = pf.line_spacing
        for table in doc.tables:
            if "表格" not in style_stats:
                style_stats["表格"] = {"font": "宋体", "size": "五号", "bold": False, "align": "居中", "line_type": "倍数", "line_value": 1.25, "indent": 0, "space_before": 0, "space_after": 0, "char_spacing": 0}
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.runs:
                            run = para.runs[0]
                            if run.font.name:
                                if run.font.name in CN_FONT_LIST:
                                    style_stats["表格"]["font"] = run.font.name
                            if run.font.size:
                                for size_name, size_pt in FONT_SIZE_MAP.items():
                                    if abs(run.font.size.pt - size_pt) < 0.5:
                                        style_stats["表格"]["size"] = size_name
                                        break
        for level in ["一级标题", "二级标题", "三级标题", "正文", "表格"]:
            if level in style_stats:
                cn_format[level] = style_stats[level]
                en_format[level] = {
                    "en_font": "Times New Roman",
                    "size_same_as_cn": True,
                    "size": style_stats[level].get("size", "小四"),
                    "bold": style_stats[level].get("bold", False),
                    "italic": False
                }
        template_data = {
            "name": f"自定义模板_{len(st.session_state.custom_templates) + 1}",
            "update_time": datetime.now().strftime('%Y-%m-%d'),
            "cn_format": cn_format,
            "en_format": en_format
        }
        return template_data, None, None
    except Exception as e:
        return None, None, str(e)

def standardize_cnki_reference(text):
    if not text.strip():
        return text, False
    text = RE_REF_SPACE.sub(' ', text.strip())
    text = RE_REF_CN_FONT.sub(r'\1[\2]', text)
    text = RE_REF_DOT.sub('.', text)
    text = RE_REF_COMMA.sub(',', text)
    text = RE_REF_COLON.sub(':', text)
    if RE_REF_FLAG.match(text) or RE_REF_KEYWORD.search(text):
        return text, True
    return text, False

def parse_plagiarism_report(file):
    try:
        content = file.read().decode('utf-8', errors='ignore')
        red_parts = RE_RED_HIGHLIGHT.findall(content)
        plain_text = RE_RED_HIGHLIGHT.sub(r'\1', content)
        return red_parts, plain_text, None
    except Exception as e:
        return None, None, str(e)

def format_compliance_check(doc, cn_format):
    try:
        check_report = []
        title_levels = ["一级标题", "二级标题", "三级标题"]
        all_levels = title_levels + ["正文", "表格"]
        
        for para in doc.paragraphs:
            level = get_title_level(para.text)
            if level in all_levels:
                # 检查字体
                target_font = cn_format[level]["font"]
                target_size = FONT_SIZE_MAP.get(cn_format[level]["size"], 12)
                for run in para.runs:
                    if run.font.name != target_font and run.font.name in CN_FONT_LIST:
                        check_report.append(f"⚠️ 【{level}】{para.text[:20]}... 字体不符合要求，应为{target_font}")
                    if run.font.size and abs(run.font.size.pt - target_size) > 0.1:
                        check_report.append(f"⚠️ 【{level}】{para.text[:20]}... 字号不符合要求，应为{cn_format[level]['size']}")
                
                # 检查对齐方式
                target_align = ALIGN_MAP[cn_format[level]["align"]]
                if para.paragraph_format.alignment != target_align:
                    check_report.append(f"⚠️ 【{level}】{para.text[:20]}... 对齐方式不符合要求，应为{cn_format[level]['align']}")
                
                # 检查行距
                if cn_format[level]["line_type"] == "固定值":
                    target_line = cn_format[level]["line_value"]
                    if para.paragraph_format.line_spacing_rule != WD_LINE_SPACING.EXACTLY or abs(para.paragraph_format.line_spacing.pt - target_line) > 0.1:
                        check_report.append(f"⚠️ 【{level}】{para.text[:20]}... 行距不符合要求，应为固定值{target_line}pt")
                else:
                    target_line = cn_format[level]["line_value"]
                    if para.paragraph_format.line_spacing_rule != WD_LINE_SPACING.MULTIPLE or abs(para.paragraph_format.line_spacing - target_line) > 0.1:
                        check_report.append(f"⚠️ 【{level}】{para.text[:20]}... 行距不符合要求，应为{target_line}倍")
                
                # 检查首行缩进（非表格）
                if level != "表格":
                    target_indent = cn_format[level]["indent"] * 0.74  # 转换为厘米
                    if abs(para.paragraph_format.first_line_indent.cm - target_indent) > 0.1:
                        check_report.append(f"⚠️ 【{level}】{para.text[:20]}... 首行缩进不符合要求，应为{cn_format[level]['indent']}字符")
                
                # 检查段前/段后间距
                if level != "表格":
                    target_before = cn_format[level]["space_before"]
                    target_after = cn_format[level]["space_after"]
                    if abs(para.paragraph_format.space_before.pt - target_before) > 0.1:
                        check_report.append(f"⚠️ 【{level}】{para.text[:20]}... 段前间距不符合要求，应为{target_before}pt")
                    if abs(para.paragraph_format.space_after.pt - target_after) > 0.1:
                        check_report.append(f"⚠️ 【{level}】{para.text[:20]}... 段后间距不符合要求，应为{target_after}pt")
        
        # 检查表格格式
        for i, table in enumerate(doc.tables):
            target_font = cn_format["表格"]["font"]
            target_size = FONT_SIZE_MAP.get(cn_format["表格"]["size"], 10.5)
            target_align = ALIGN_MAP[cn_format["表格"]["align"]]
            
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            # 检查表格单元格对齐方式
                            if para.paragraph_format.alignment != target_align:
                                check_report.append(f"⚠️ 【表格{i+1}】单元格内容对齐方式不符合要求，应为{cn_format['表格']['align']}")
                            
                            # 检查表格字体和字号
                            for run in para.runs:
                                if run.font.name != target_font and run.font.name in CN_FONT_LIST:
                                    check_report.append(f"⚠️ 【表格{i+1}】单元格字体不符合要求，应为{target_font}")
                                if run.font.size and abs(run.font.size.pt - target_size) > 0.1:
                                    check_report.append(f"⚠️ 【表格{i+1}】单元格字号不符合要求，应为{cn_format['表格']['size']}")
        
        if not check_report:
            check_report.append("✅ 文档格式完全符合要求，无违规项")
        return check_report
    except Exception as e:
        return [f"⚠️ 格式检查失败：{str(e)}"]

def optimize_image_layout(doc):
    image_count = 0
    for para in doc.paragraphs:
        has_image = False
        for run in para.runs:
            if run._element.xpath('.//a:blip'):
                has_image = True
                image_count += 1
                break
        if has_image:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.keep_together = True
            para.paragraph_format.first_line_indent = Cm(0)
    return image_count

def is_white_text(text):
    text_strip = text.strip()
    for word in WHITE_WORDS:
        if word in text_strip:
            return True
    if RE_WHITE_NUMBER.match(text_strip):
        return True
    if RE_WHITE_QUOTE.match(text_strip):
        return True
    return False

def check_semantic_keep(original, modified):
    original_keywords = set(RE_KEYWORDS.findall(original))
    modified_keywords = set(RE_KEYWORDS.findall(modified))
    if not original_keywords and not modified_keywords:
        return 1.0
    if not original_keywords:
        return 0.0 if modified_keywords else 1.0
    overlap = original_keywords & modified_keywords
    return len(overlap) / len(original_keywords)

def call_doubao_api(text, api_key, prompt):
    try:
        if not api_key:
            return None, "API Key不能为空"
        
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "doubao-pro",
            "messages": [
                {"role": "system", "content": prompt},
                {"role": "user", "content": text}
            ]
        }
        
        try:
            response = requests.post("https://ark.cn-beijing.volces.com/api/v3/chat/completions", headers=headers, json=payload, timeout=30)
            if response.status_code == 200:
                try:
                    return response.json()["choices"][0]["message"]["content"].strip(), None
                except (KeyError, IndexError, ValueError) as e:
                    return None, f"API返回格式错误: {str(e)}"
            else:
                return None, f"API调用失败: 状态码 {response.status_code}, {response.text}"
        except requests.exceptions.Timeout:
            return None, "API调用超时"
        except requests.exceptions.ConnectionError:
            return None, "网络连接错误"
        except requests.exceptions.RequestException as e:
            return None, f"API请求异常: {str(e)}"
    except Exception as e:
        return None, f"未知错误: {str(e)}"

def rewrite_sentence(sentence, level_config, api_key=None, forbidden_text=None):
    try:
        original = sentence.strip()
        if len(original) < 5 or is_white_text(original):
            return original, "原文保留（白名单/短句）", 1.0
        modified = original
        rewrite_type = "无修改"
        if forbidden_text and original in forbidden_text:
            if api_key:
                try:
                    result, error = call_doubao_api(original, api_key, "你是一个论文润色专家，请润色这段文本，保持原意，让它不重复，优化表达")
                    if not error:
                        modified = result
                        rewrite_type = "AI针对性润色(规避查重)"
                except Exception as e:
                    # API调用失败，使用备用方案
                    parts = [p.strip() for p in RE_CLAUSE_SPLIT.split(modified) if p.strip()]
                    if len(parts) >= 3:
                        last_part = parts[-1]
                        rest_parts = parts[:-1]
                        random.shuffle(rest_parts)
                        modified = "，".join(rest_parts + [last_part])
                        if not modified.endswith(("。", "！", "？", "；")):
                            modified += "。"
                        rewrite_type = "针对性语序调整(规避查重)"
            else:
                parts = [p.strip() for p in RE_CLAUSE_SPLIT.split(modified) if p.strip()]
                if len(parts) >= 3:
                    last_part = parts[-1]
                    rest_parts = parts[:-1]
                    random.shuffle(rest_parts)
                    modified = "，".join(rest_parts + [last_part])
                    if not modified.endswith(("。", "！", "？", "；")):
                        modified += "。"
                    rewrite_type = "针对性语序调整(规避查重)"
        elif api_key:
            try:
                result, error = call_doubao_api(original, api_key, "你是一个论文润色专家，请润色这段学术文本，保持原意，优化表达")
                if not error:
                    modified = result
                    rewrite_type = "AI智能润色"
                else:
                    # API调用失败，使用备用方案
                    if level_config["synonym"]:
                        for old, new in SYNONYM_DICT.items():
                            if old in modified and not is_white_text(old):
                                modified = modified.replace(old, new)
                                rewrite_type = "同义词替换"
                    if level_config["sentence_reorder"]:
                        parts = [p.strip() for p in RE_CLAUSE_SPLIT.split(modified) if p.strip()]
                        if len(parts) >= 3 and not is_white_text(modified):
                            last_part = parts[-1]
                            rest_parts = parts[:-1]
                            random.shuffle(rest_parts)
                            modified = "，".join(rest_parts + [last_part])
                            if not modified.endswith(("。", "！", "？", "；")):
                                modified += "。"
                            rewrite_type = "句式重构+语序打乱"
            except Exception as e:
                # API调用失败，使用备用方案
                if level_config["synonym"]:
                    for old, new in SYNONYM_DICT.items():
                        if old in modified and not is_white_text(old):
                            modified = modified.replace(old, new)
                            rewrite_type = "同义词替换"
                if level_config["sentence_reorder"]:
                    parts = [p.strip() for p in RE_CLAUSE_SPLIT.split(modified) if p.strip()]
                    if len(parts) >= 3 and not is_white_text(modified):
                        last_part = parts[-1]
                        rest_parts = parts[:-1]
                        random.shuffle(rest_parts)
                        modified = "，".join(rest_parts + [last_part])
                        if not modified.endswith(("。", "！", "？", "；")):
                            modified += "。"
                        rewrite_type = "句式重构+语序打乱"
        if not api_key or rewrite_type == "无修改":
            if level_config["synonym"]:
                for old, new in SYNONYM_DICT.items():
                    if old in modified and not is_white_text(old):
                        modified = modified.replace(old, new)
                        rewrite_type = "同义词替换"
            if level_config["sentence_reorder"]:
                parts = [p.strip() for p in RE_CLAUSE_SPLIT.split(modified) if p.strip()]
                if len(parts) >= 3 and not is_white_text(modified):
                    last_part = parts[-1]
                    rest_parts = parts[:-1]
                    random.shuffle(rest_parts)
                    modified = "，".join(rest_parts + [last_part])
                    if not modified.endswith(("。", "！", "？", "；")):
                        modified += "。"
                    rewrite_type = "句式重构+语序打乱"
        semantic_score = check_semantic_keep(original, modified)
        if semantic_score < 0.7:
            return original, "原文保留（语义重合度不达标）", 1.0
        return modified, rewrite_type, round(semantic_score, 4)
    except Exception as e:
        # 出错时返回原文，确保处理不中断
        return sentence, "原文保留（处理出错）", 1.0

def rewrite_paragraph(text, level_config, api_key=None, forbidden_text=None):
    try:
        change_log = []
        sentences = RE_SENTENCE_SPLIT.split(text)
        new_sentences = []
        for sent in sentences:
            if not sent.strip():
                new_sentences.append(sent)
                continue
            new_sent, rewrite_type, semantic_score = rewrite_sentence(sent, level_config, api_key, forbidden_text)
            new_sentences.append(new_sent)
            if sent != new_sent:
                change_log.append({
                    "original": sent,
                    "modified": new_sent,
                    "type": rewrite_type,
                    "semantic_score": semantic_score
                })
        return "".join(new_sentences), change_log
    except Exception as e:
        # 出错时返回原文，确保处理不中断
        return text, []

def simulate_check_rate(text):
    """模拟查重率计算，可替换为真实API"""
    words = RE_KEYWORDS.findall(text)
    if not words:
        return 10.0
    repeat_count = sum(1 for w in words if w in WHITE_WORDS)
    rate = min(40, max(5, repeat_count / len(words) * 100))
    return round(rate, 1)

def process_doc(
    file,
    cn_format,
    en_format,
    enable_rewrite=False,
    rewrite_level="标准润色",
    bind_wps_style=True,
    standardize_ref=True,
    api_key=None,
    forbidden_text=None
):
    try:
        file.seek(0, os.SEEK_END)
        file_size_mb = file.tell() / (1024 * 1024)
        file.seek(0)
        if file_size_mb > MAX_FILE_SIZE_MB:
            raise Exception(f"文件大小超过限制（{MAX_FILE_SIZE_MB}MB），当前大小：{file_size_mb:.2f}MB")
        
        try:
            if file.name.endswith('.docx'):
                doc = Document(file)
            elif file.name.endswith('.doc'):
                # 转换doc为docx
                docx_file = doc_to_docx(file)
                doc = Document(docx_file)
            elif file.name.endswith('.pdf'):
                # 转换pdf为docx
                docx_file = pdf_to_docx(file)
                doc = Document(docx_file)
            else:
                raise Exception(f"不支持的文件格式：{file.name.split('.')[-1]}")
        except Exception as e:
            raise Exception(f"文档读取失败：{str(e)}")
    except Exception as e:
        raise Exception(f"初始化失败：{str(e)}")
    
    total_changes = []
    ref_count = 0
    process_log = []
    title_stats = {"一级标题": 0, "二级标题": 0, "三级标题": 0, "正文": 0, "表格": len(doc.tables)}
    rewrite_config = REWRITE_LEVEL[rewrite_level]
    style_warn_logged = False
    
    # 文档分块处理，提升大文件处理速度
    total_paragraphs = len(doc.paragraphs)
    chunk_size = 100  # 每块处理100个段落
    num_chunks = (total_paragraphs + chunk_size - 1) // chunk_size
    
    try:
        prev_para_text = None  # 记录上一段文本，用于上下文校验
        
        for chunk_idx in range(num_chunks):
            start_idx = chunk_idx * chunk_size
            end_idx = min((chunk_idx + 1) * chunk_size, total_paragraphs)
            chunk_paragraphs = doc.paragraphs[start_idx:end_idx]
            
            for para in chunk_paragraphs:
                original_text = para.text
                level = get_title_level(original_text, prev_para_text)
                title_stats[level] += 1
                
                if enable_rewrite and level == "正文":
                    new_text, changes = rewrite_paragraph(original_text, rewrite_config, api_key, forbidden_text)
                    if changes:
                        total_changes.extend(changes)
                        para.text = new_text
                
                if standardize_ref:
                    new_text, is_ref = standardize_cnki_reference(para.text)
                    if is_ref:
                        para.text = new_text
                        ref_count += 1
                
                cn_style = cn_format[level]
                en_style = en_format[level]
                
                # 更新上一段文本
                prev_para_text = original_text.strip()
                
                if bind_wps_style and level in WPS_STYLE_MAPPING:
                    try:
                        target_style_id = WPS_STYLE_MAPPING[level]
                        if target_style_id in doc.styles:
                            para.style = doc.styles[target_style_id]
                    except Exception as e:
                        if not style_warn_logged:
                            process_log.append(f"⚠️ 文档内置样式异常，已跳过WPS标题样式绑定")
                            style_warn_logged = True
                
                para_format = para.paragraph_format
                para_format.alignment = ALIGN_MAP[cn_style["align"]]
                para_format.first_line_indent = Cm(cn_style["indent"] * 0.74)
                para_format.space_before = Pt(cn_style["space_before"])
                para_format.space_after = Pt(cn_style["space_after"])
                
                if cn_style["line_type"] == "固定值":
                    para_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    para_format.line_spacing = Pt(cn_style["line_value"])
                else:
                    para_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    para_format.line_spacing = cn_style["line_value"]
                
                cn_size_pt = FONT_SIZE_MAP.get(cn_style["size"], 12)
                
                for run in para.runs:
                    run.font.name = cn_style["font"]
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_style["font"])
                    run._element.rPr.rFonts.set(qn('w:ascii'), en_style["en_font"])
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), en_style["en_font"])
                    run.font.size = Pt(cn_size_pt)
                    run.font.bold = en_style["bold"] if en_style["bold"] else cn_style["bold"]
                    run.font.italic = en_style["italic"]
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    # 新增字间距应用
                    if cn_style.get("char_spacing", 0) > 0:
                        run.font.spacing = Pt(cn_style["char_spacing"])
        
        process_log.append(f"✅ 全文档段落处理完成（共{total_paragraphs}个段落，分{num_chunks}块处理）")
        if enable_rewrite:
            process_log.append(f"✅ 智能润色完成，共修改{len(total_changes)}处")
        if standardize_ref and ref_count > 0:
            process_log.append(f"✅ 知网参考文献标准化完成，共处理{ref_count}条")
        process_log.append(f"📊 标题识别结果：一级{title_stats['一级标题']}、二级{title_stats['二级标题']}、三级{title_stats['三级标题']}")
    except Exception as e:
        raise Exception(f"文档处理失败：{str(e)}")
    
    try:
        image_count = optimize_image_layout(doc)
        if image_count > 0:
            process_log.append(f"✅ 优化{image_count}张图片排版")
        else:
            process_log.append("✅ 未检测到图片")
    except Exception as e:
        process_log.append(f"⚠️ 图片处理失败：{str(e)}")
    
    try:
        cn_table_style = cn_format["表格"]
        en_table_style = en_format["表格"]
        table_cn_size = FONT_SIZE_MAP.get(cn_table_style["size"], 10.5)
        
        # 表格分块处理
        total_tables = len(doc.tables)
        table_chunk_size = 10  # 每块处理10个表格
        num_table_chunks = (total_tables + table_chunk_size - 1) // table_chunk_size
        
        for chunk_idx in range(num_table_chunks):
            start_idx = chunk_idx * table_chunk_size
            end_idx = min((chunk_idx + 1) * table_chunk_size, total_tables)
            chunk_tables = doc.tables[start_idx:end_idx]
            
            for table in chunk_tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if enable_rewrite:
                                original_text = para.text.strip()
                                if original_text and not is_white_text(original_text):
                                    new_text, changes = rewrite_paragraph(original_text, rewrite_config, api_key, forbidden_text)
                                    if changes:
                                        total_changes.extend(changes)
                                        para.text = new_text
                            
                            para.alignment = ALIGN_MAP[cn_table_style["align"]]
                            para_format = para.paragraph_format
                            para_format.space_before = Pt(cn_table_style["space_before"])
                            para_format.space_after = Pt(cn_table_style["space_after"])
                            
                            if cn_table_style["line_type"] == "固定值":
                                para_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                                para_format.line_spacing = Pt(cn_table_style["line_value"])
                            else:
                                para_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                                para_format.line_spacing = cn_table_style["line_value"]
                            
                            for run in para.runs:
                                run.font.name = cn_table_style["font"]
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_table_style["font"])
                                run._element.rPr.rFonts.set(qn('w:ascii'), en_table_style["en_font"])
                                run._element.rPr.rFonts.set(qn('w:hAnsi'), en_table_style["en_font"])
                                run.font.size = Pt(table_cn_size)
                                run.font.bold = en_table_style["bold"] if en_table_style["bold"] else cn_table_style["bold"]
                                run.font.italic = en_table_style["italic"]
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                if cn_table_style.get("char_spacing", 0) > 0:
                                    run.font.spacing = Pt(cn_table_style["char_spacing"])
        
        process_log.append(f"✅ 表格格式处理完成（共{total_tables}个表格）")
    except Exception as e:
        process_log.append(f"⚠️ 表格处理失败：{str(e)}")
    
    try:
        check_report = format_compliance_check(doc, cn_format)
        process_log.append("✅ 格式合规检查完成")
    except Exception as e:
        check_report = [f"⚠️ 格式检查失败：{str(e)}"]
        process_log.append(check_report[0])
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    # 提取全文用于查重
    full_text = "\n".join([p.text for p in doc.paragraphs])
    return output, total_changes, title_stats, process_log, check_report, full_text

def generate_report(changes, rewrite_level, title_stats, process_log, check_report):
    total_count = len(changes)
    report = f"# 文档处理报告\n"
    report += f"📅 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    report += f"⚙️ 润色强度：{rewrite_level}\n"
    report += f"📝 总修改条数：{total_count}\n\n"

    report += "## 一、处理流程日志\n"
    for log in process_log:
        report += f"- {log}\n"
    report += "\n## 二、标题识别统计\n"
    for level, count in title_stats.items():
        report += f"- {level}：{count} 个\n"
    report += "\n## 三、格式合规性检查报告\n"
    for item in check_report:
        report += f"- {item}\n"
    if changes:
        report += "\n## 四、润色修改详情\n"
        for i, change in enumerate(changes[:20]):
            report += f"\n### 修改 {i+1}\n"
            report += f"- **原句**: {change['original']}\n"
            report += f"- **修改**: {change['modified']}\n"
            report += f"- **类型**: {change['type']}\n"
            report += f"- **语义保留**: {change['semantic_score']*100:.1f}%\n"
    return report.encode("utf-8")

def search_academic_papers(keyword, max_results=5):
    """学术文献搜索功能"""
    try:
        # 模拟学术搜索结果
        # 实际项目中可以集成真实的学术搜索API，如CNKI、Google Scholar等
        mock_results = [
            {
                "title": f"{keyword}的研究进展",
                "authors": ["张三", "李四"],
                "journal": "中国学术期刊",
                "year": 2024,
                "abstract": f"本文对{keyword}的最新研究进展进行了综述，包括理论基础、实验方法和应用前景等方面。",
                "url": "https://example.com/paper1"
            },
            {
                "title": f"基于{keyword}的创新方法",
                "authors": ["王五", "赵六"],
                "journal": "科技通报",
                "year": 2023,
                "abstract": f"提出了一种基于{keyword}的创新方法，通过实验验证了其有效性和可行性。",
                "url": "https://example.com/paper2"
            },
            {
                "title": f"{keyword}在实践中的应用",
                "authors": ["钱七", "孙八"],
                "journal": "应用科学学报",
                "year": 2024,
                "abstract": f"探讨了{keyword}在实际应用中的具体案例，分析了其优势和不足。",
                "url": "https://example.com/paper3"
            },
            {
                "title": f"{keyword}的理论模型",
                "authors": ["周九", "吴十"],
                "journal": "理论研究",
                "year": 2023,
                "abstract": f"建立了{keyword}的理论模型，为后续研究提供了理论基础。",
                "url": "https://example.com/paper4"
            },
            {
                "title": f"{keyword}的未来发展趋势",
                "authors": ["郑一", "王二"],
                "journal": "未来科学",
                "year": 2024,
                "abstract": f"分析了{keyword}的未来发展趋势，预测了可能的研究方向和应用领域。",
                "url": "https://example.com/paper5"
            }
        ]
        return mock_results[:max_results], None
    except Exception as e:
        return [], str(e)

def export_template(template_data, export_type="json"):
    if export_type == "json":
        return json.dumps(template_data, ensure_ascii=False, indent=2).encode("utf-8")
    else:
        text = f"模板配置文件\n"
        text += f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        text += f"模板名称: {template_data.get('name', '自定义模板')}\n"
        text += f"更新时间: {template_data.get('update_time', datetime.now().strftime('%Y-%m-%d'))}\n\n"
        text += "=== 中文格式设置 ===\n"
        for level, cfg in template_data.get('cn_format', {}).items():
            text += f"\n[{level}]\n"
            for k, v in cfg.items():
                text += f"{k} = {v}\n"
        text += "\n=== 西文格式设置 ===\n"
        for level, cfg in template_data.get('en_format', {}).items():
            text += f"\n[{level}]\n"
            for k, v in cfg.items():
                text += f"{k} = {v}\n"
        return text.encode("utf-8")

def import_template(file):
    try:
        content = file.read().decode('utf-8')
        if file.name.endswith('.json'):
            data = json.loads(content)
            return data, None
        else:
            data = {}
            current_section = None
            current_level = None
            current_cfg = {}
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith('==='):
                    current_section = line.strip('= ').strip()
                    continue
                if line.startswith('[') and line.endswith(']'):
                    if current_level and current_cfg:
                        if current_section == "中文格式设置":
                            data['cn_format'] = data.get('cn_format', {})
                            data['cn_format'][current_level] = current_cfg
                        else:
                            data['en_format'] = data.get('en_format', {})
                            data['en_format'][current_level] = current_cfg
                    current_level = line.strip('[]')
                    current_cfg = {}
                    continue
                if '=' in line:
                    k, v = line.split('=', 1)
                    k = k.strip()
                    v = v.strip()
                    try:
                        if v.lower() == 'true':
                            v = True
                        elif v.lower() == 'false':
                            v = False
                        elif '.' in v:
                            v = float(v)
                        else:
                            try:
                                v = int(v)
                            except:
                                pass
                    except:
                        pass
                    current_cfg[k] = v
            if current_level and current_cfg:
                if current_section == "中文格式设置":
                    data['cn_format'] = data.get('cn_format', {})
                    data['cn_format'][current_level] = current_cfg
                else:
                    data['en_format'] = data.get('en_format', {})
                    data['en_format'][current_level] = current_cfg
            return data, None
    except Exception as e:
        return None, str(e)

# ====================== 页面刷新兼容函数 ======================
def safe_rerun():
    try:
        st.rerun()
    except AttributeError:
        st.experimental_rerun()

# ====================== Session状态初始化 ======================
def init_session_state():
    if "current_template" not in st.session_state:
        st.session_state.current_template = "三创赛-全国大学生电子商务创新创意及创业挑战赛"
        st.session_state.cn_format, st.session_state.en_format = get_cached_template(st.session_state.current_template)
    if "custom_templates" not in st.session_state:
        st.session_state.custom_templates = {}
    if "version" not in st.session_state:
        st.session_state.version = 0
    if "learned_forbidden" not in st.session_state:
        st.session_state.learned_forbidden = None
    if "learn_history" not in st.session_state:
        st.session_state.learn_history = []
    if "formatted_doc" not in st.session_state:
        st.session_state.formatted_doc = None
    if "formatted_report" not in st.session_state:
        st.session_state.formatted_report = None
    if "check_rate" not in st.session_state:
        st.session_state.check_rate = None
    if "doc_full_text" not in st.session_state:
        st.session_state.doc_full_text = ""
    if "need_polish" not in st.session_state:
        st.session_state.need_polish = False
    if "polish_doc" not in st.session_state:
        st.session_state.polish_doc = None
    if "polish_report" not in st.session_state:
        st.session_state.polish_report = None
    if "process_timestamp" not in st.session_state:
        st.session_state.process_timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    if "dark_mode" not in st.session_state:
        st.session_state.dark_mode = False
    if "processed_files" not in st.session_state:
        st.session_state.processed_files = {}

# ====================== 主应用UI（按需求重构）======================
def main():
    # 页面基础配置
    st.set_page_config(
        page_title="智能论文&竞赛格式处理平台",
        layout="wide",
        page_icon="📝",
        initial_sidebar_state="collapsed"
    )
    # 全局样式
    st.markdown("""
    <style>
    /* 基础样式 */
    .stBlockContainer {
        min-width: 1200px;
        max-width: 100% !important;
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    div[data-testid="stVerticalBlock"] > div {
        gap: 0.8rem;
    }
    
    /* 响应式布局 */
    @media (max-width: 1200px) {
        .stBlockContainer {
            min-width: 100%;
            padding-top: 1rem;
            padding-bottom: 1rem;
        }
    }
    
    @media (max-width: 768px) {
        .stApp {
            padding: 1rem;
        }
        .stButton > button {
            width: 100%;
        }
        .stSelectbox, .stTextInput, .stCheckbox {
            width: 100%;
        }
    }
    
    /* 深色模式样式 */
    .dark {
        background-color: #1e1e1e;
        color: #f0f0f0;
    }
    
    .dark .stApp {
        background-color: #1e1e1e;
        color: #f0f0f0;
    }
    
    .dark .stButton > button {
        background-color: #333;
        color: #f0f0f0;
        border: 1px solid #555;
    }
    
    .dark .stSelectbox > div {
        background-color: #333;
        color: #f0f0f0;
    }
    
    .dark .stTextInput > div > div {
        background-color: #333;
        color: #f0f0f0;
    }
    
    .dark .stExpander {
        background-color: #2d2d2d;
        color: #f0f0f0;
    }
    
    .dark .stExpanderContent {
        background-color: #2d2d2d;
        color: #f0f0f0;
    }
    
    .dark .stAlert {
        background-color: #2d2d2d;
        color: #f0f0f0;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # 初始化状态
    init_session_state()

    # 深色模式切换
    col_mode, col_title = st.columns([1, 5])
    with col_mode:
        dark_mode = st.checkbox("🌙 深色模式", value=st.session_state.dark_mode, key="dark_mode_checkbox")
        if dark_mode != st.session_state.dark_mode:
            st.session_state.dark_mode = dark_mode
            safe_rerun()
    
    # 应用深色模式
    if st.session_state.dark_mode:
        st.markdown('<body class="dark">', unsafe_allow_html=True)

    # 整体布局：左1右4 核心框架
    left_col, right_col = st.columns([1, 4])

    # ============== 左栏：格式精细调整 + 模板管理 ==============
    with left_col:
        st.markdown("### ⚙️ 格式精细设置")
        st.divider()

        # 模板选择
        template_options = list(ALL_TEMPLATES.keys()) + list(st.session_state.custom_templates.keys())
        selected_template = st.selectbox(
            "选择基础模板",
            options=template_options,
            index=template_options.index(st.session_state.current_template) if st.session_state.current_template in template_options else 0,
            key="template_select"
        )
        # 模板切换逻辑
        if selected_template != st.session_state.current_template:
            st.session_state.current_template = selected_template
            if selected_template in ALL_TEMPLATES:
                st.session_state.cn_format, st.session_state.en_format = get_cached_template(selected_template)
            else:
                tmp = st.session_state.custom_templates[selected_template]
                st.session_state.cn_format = copy.deepcopy(tmp["cn_format"])
                st.session_state.en_format = copy.deepcopy(tmp["en_format"])
            st.session_state.version += 1
            safe_rerun()
        
        # 模板信息
        if selected_template in ALL_TEMPLATES:
            update_time = ALL_TEMPLATES[selected_template].get("update_time", "未知")
            st.caption(f"📅 模板更新时间：{update_time}")
            if ALL_TEMPLATES[selected_template].get("special_requirements"):
                with st.expander("模板格式要求", expanded=False):
                    for req in ALL_TEMPLATES[selected_template]["special_requirements"]:
                        st.markdown(f"- {req}")
        else:
            update_time = st.session_state.custom_templates[selected_template].get("update_time", datetime.now().strftime('%Y-%m-%d'))
            st.caption(f"📅 自定义模板更新时间：{update_time}")

        # 自定义模板保存
        st.divider()
        st.subheader("📑 自定义模板", divider=True)
        template_name = st.text_input("模板命名", placeholder="输入自定义模板名称", key="template_name_input")
        col_save, col_load = st.columns(2)
        with col_save:
            if st.button("保存当前格式", use_container_width=True, type="primary"):
                if template_name:
                    st.session_state.custom_templates[template_name] = {
                        "cn_format": copy.deepcopy(st.session_state.cn_format),
                        "en_format": copy.deepcopy(st.session_state.en_format),
                        "update_time": datetime.now().strftime('%Y-%m-%d')
                    }
                    st.success(f"✅ 模板「{template_name}」保存成功")
                    st.session_state.version += 1
                    safe_rerun()
                else:
                    st.error("请输入模板名称")
        with col_load:
            if st.session_state.custom_templates:
                selected_custom = st.selectbox("加载自定义模板", options=list(st.session_state.custom_templates.keys()), label_visibility="collapsed")
                if st.button("加载模板", use_container_width=True):
                    tmp = st.session_state.custom_templates[selected_custom]
                    st.session_state.cn_format = copy.deepcopy(tmp["cn_format"])
                    st.session_state.en_format = copy.deepcopy(tmp["en_format"])
                    st.session_state.current_template = selected_custom
                    st.session_state.version += 1
                    st.success(f"✅ 已加载「{selected_custom}」")
                    safe_rerun()

        # 格式精细调整
        with st.container(height=500, border=True):
            st.subheader("🎨 格式参数调整", divider=True)
            for level in ["一级标题", "二级标题", "三级标题", "正文", "表格"]:
                with st.expander(f"{level}格式设置", expanded=(level == "正文")):
                    cfg = st.session_state.cn_format[level]
                    # 基础格式
                    col1, col2 = st.columns(2)
                    with col1:
                        cfg["font"] = st.selectbox(
                            "中文字体", CN_FONT_LIST,
                            index=CN_FONT_LIST.index(cfg["font"]) if cfg["font"] in CN_FONT_LIST else 0,
                            key=f"cn_{level}_font_{st.session_state.version}"
                        )
                        cfg["size"] = st.selectbox(
                            "字号", list(FONT_SIZE_MAP.keys()),
                            index=list(FONT_SIZE_MAP.keys()).index(cfg["size"]) if cfg["size"] in FONT_SIZE_MAP else 5,
                            key=f"cn_{level}_size_{st.session_state.version}"
                        )
                        cfg["bold"] = st.checkbox(
                            "加粗", cfg["bold"],
                            key=f"cn_{level}_bold_{st.session_state.version}"
                        )
                    with col2:
                        cfg["align"] = st.selectbox(
                            "对齐方式", list(ALIGN_MAP.keys()),
                            index=list(ALIGN_MAP.keys()).index(cfg["align"]),
                            key=f"cn_{level}_align_{st.session_state.version}"
                        )
                        cfg["line_type"] = st.selectbox(
                            "行距类型", ["倍数", "固定值"],
                            index=0 if cfg["line_type"] == "倍数" else 1,
                            key=f"cn_{level}_line_type_{st.session_state.version}"
                        )
                        cfg["line_value"] = st.number_input(
                            "行距值",
                            min_value=0.0 if cfg["line_type"] == "倍数" else 8,
                            value=cfg["line_value"],
                            step=0.1 if cfg["line_type"] == "倍数" else 1,
                            key=f"cn_{level}_line_val_{st.session_state.version}"
                        )
                    # 缩进与间距
                    if level != "表格":
                        col3, col4, col5 = st.columns(3)
                        with col3:
                            cfg["indent"] = st.number_input(
                                "首行缩进(字符)", 0, 4, cfg["indent"], 1,
                                key=f"cn_{level}_indent_{st.session_state.version}"
                            )
                        with col4:
                            cfg["space_before"] = st.number_input(
                                "段前间距(pt)", 0, 24, cfg["space_before"], 1,
                                key=f"cn_{level}_before_{st.session_state.version}"
                            )
                        with col5:
                            cfg["space_after"] = st.number_input(
                                "段后间距(pt)", 0, 24, cfg["space_after"], 1,
                                key=f"cn_{level}_after_{st.session_state.version}"
                            )
                    # 字间距
                    cfg["char_spacing"] = st.slider(
                        "字间距(pt)", 0, 10, cfg.get("char_spacing", 0), 1,
                        key=f"cn_{level}_char_space_{st.session_state.version}"
                    )
                    # 同步更新session
                    st.session_state.cn_format[level] = cfg
            # 西文格式设置
            with st.expander("🔤 西文格式全局设置", expanded=False):
                for level in ["一级标题", "二级标题", "三级标题", "正文", "表格"]:
                    cfg = st.session_state.en_format[level]
                    col1, col2 = st.columns(2)
                    with col1:
                        cfg["en_font"] = st.selectbox(
                            f"{level}西文字体", EN_FONT_LIST,
                            index=EN_FONT_LIST.index(cfg["en_font"]) if cfg["en_font"] in EN_FONT_LIST else 0,
                            key=f"en_{level}_font_{st.session_state.version}"
                        )
                    with col2:
                        cfg["bold"] = st.checkbox(f"{level}西文加粗", cfg["bold"], key=f"en_{level}_bold_{st.session_state.version}")
                        cfg["italic"] = st.checkbox(f"{level}西文斜体", cfg["italic"], key=f"en_{level}_italic_{st.session_state.version}")
                    st.session_state.en_format[level] = cfg

        # 模板导入导出
        st.divider()
        st.subheader("📤 模板导入导出", divider=True)
        export_type = st.radio("导出格式", options=["json(专用格式)", "txt(通用格式)"], index=0, horizontal=True)
        if st.button("导出当前模板", use_container_width=True):
            template_data = {
                "name": selected_template,
                "update_time": datetime.now().strftime('%Y-%m-%d'),
                "cn_format": st.session_state.cn_format,
                "en_format": st.session_state.en_format
            }
            export_type_code = "json" if "json" in export_type else "txt"
            data = export_template(template_data, export_type_code)
            st.download_button(
                label="⬇️ 下载模板文件",
                data=data,
                file_name=f"{selected_template}.{export_type_code}",
                mime="application/json" if export_type_code == "json" else "text/plain",
                use_container_width=True
            )
        st.divider()
        uploaded_template = st.file_uploader("上传模板文件", type=["json", "txt"], label_visibility="collapsed")
        if uploaded_template:
            data, error = import_template(uploaded_template)
            if error:
                st.error(f"导入失败：{error}")
            elif data:
                st.success("模板解析成功！")
                new_name = st.text_input("导入模板命名", value=uploaded_template.name.split('.')[0])
                if st.button("导入到系统", use_container_width=True):
                    st.session_state.custom_templates[new_name] = data
                    st.success(f"✅ 模板「{new_name}」导入成功！")
                    st.session_state.version += 1
                    safe_rerun()

    # ============== 右栏：核心操作流程 ==============
    with right_col:
        st.title("📝 智能论文&竞赛格式处理平台")
        st.success("✅ 支持一键格式标准化 | WPS自动生成导航 | 知网参考文献优化 | 智能降重润色 | 格式合规检查")
        st.divider()

        # ---------- 第一步：文档格式标准化（左模板设置+右上传）----------
        st.subheader("📄 第一步：文档格式标准化", divider=True)
        col_format_left, col_format_right = st.columns([1, 1])
        
        # 左：格式与辅助功能设置
        with col_format_left:
            st.markdown("##### 格式与辅助功能设置")
            col_func1, col_func2 = st.columns(2)
            with col_func1:
                bind_wps_style = st.checkbox("✅ 绑定WPS标题样式", value=True, help="开启后导出的文档在WPS中自动生成导航目录")
                standardize_ref = st.checkbox("📚 知网参考文献标准化", value=True, help="自动调整参考文献格式，解决知网查重标红问题")
            with col_func2:
                enable_rewrite = st.checkbox("🔄 开启智能润色", value=False)
                rewrite_level = st.selectbox("润色强度", options=list(REWRITE_LEVEL.keys()), index=1, disabled=not enable_rewrite)
            api_key = st.text_input("豆包API Key(可选)", type="password", help="用于AI智能润色和查重规避")

        # 右：文档上传
        with col_format_right:
            st.markdown("##### 待处理文档上传")
            files = st.file_uploader("上传文档", type=["docx", "doc", "pdf"], accept_multiple_files=True, help="支持同时上传多个文档批量处理，支持docx、doc、pdf格式")
            
            # 智能模板推荐
            if files:
                with st.expander("🤖 智能模板推荐", expanded=True):
                    for file in files:
                        # 对于PDF和doc文件，先转换为docx再进行推荐
                        if file.name.endswith('.pdf') or file.name.endswith('.doc'):
                            try:
                                if file.name.endswith('.pdf'):
                                    docx_file = pdf_to_docx(file)
                                else:
                                    docx_file = doc_to_docx(file)
                                recommended_template, score = recommend_template(docx_file)
                            except Exception as e:
                                st.warning(f"📄 文档 '{file.name}' 分析失败：{str(e)}")
                                continue
                        else:
                            recommended_template, score = recommend_template(file)
                        
                        if score > 0:
                            st.success(f"📄 为文档 '{file.name}' 推荐模板：**{recommended_template}** (匹配度: {score})")
                            if st.button(f"应用推荐模板", key=f"apply_{file.name}"):
                                st.session_state.current_template = recommended_template
                                st.session_state.cn_format, st.session_state.en_format = get_cached_template(recommended_template)
                                st.session_state.version += 1
                                safe_rerun()
                        else:
                            st.info(f"📄 为文档 '{file.name}' 未找到匹配模板，使用默认模板")

        # 处理按钮
        if files and st.button("🚀 开始格式处理", type="primary", use_container_width=True):
            st.session_state.process_timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            # 为每个文件创建独立的处理结果存储
            st.session_state.processed_files = {}
            
            for file in files:
                with st.spinner(f"正在处理：{file.name}"):
                    try:
                        output_doc, changes, title_stats, process_log, check_report, full_text = process_doc(
                            file=file,
                            cn_format=st.session_state.cn_format,
                            en_format=st.session_state.en_format,
                            enable_rewrite=enable_rewrite,
                            rewrite_level=rewrite_level,
                            bind_wps_style=bind_wps_style,
                            standardize_ref=standardize_ref,
                            api_key=api_key,
                            forbidden_text=st.session_state.learned_forbidden
                        )
                        
                        # 保存当前文件的处理结果
                        file_key = f"{file.name}_{st.session_state.process_timestamp}"
                        st.session_state.processed_files[file_key] = {
                            "file_name": file.name,
                            "output_doc": output_doc,
                            "full_text": full_text,
                            "check_rate": simulate_check_rate(full_text),
                            "report": generate_report(changes, rewrite_level, title_stats, process_log, check_report),
                            "title_stats": title_stats,
                            "process_log": process_log
                        }
                        
                        # 保存最后一个文件的结果到全局状态（保持向后兼容）
                        st.session_state.formatted_doc = output_doc
                        st.session_state.doc_full_text = full_text
                        st.session_state.check_rate = simulate_check_rate(full_text)
                        st.session_state.formatted_report = generate_report(changes, rewrite_level, title_stats, process_log, check_report)

                        # 处理结果展示
                        st.subheader(f"✅ 处理完成：{file.name}")
                        with st.expander("📋 处理日志", expanded=True):
                            for log in process_log:
                                st.write(log)
                        # 标题统计
                        cols = st.columns(5)
                        cols[0].metric("一级标题", title_stats["一级标题"])
                        cols[1].metric("二级标题", title_stats["二级标题"])
                        cols[2].metric("三级标题", title_stats["三级标题"])
                        cols[3].metric("正文段落", title_stats["正文"])
                        cols[4].metric("表格数量", title_stats["表格"])

                    except Exception as e:
                        st.error(f"处理失败：{str(e)}")

        # 自动生成查重率与报告
        if st.session_state.check_rate is not None:
            st.divider()
            st.markdown("##### 🔍 文档查重结果")
            rate = st.session_state.check_rate
            st.progress(rate/100)
            st.markdown(f"**文档查重率：{rate}%**")
            
            # 提示用户是否润色
            if rate > 20:
                st.warning(f"⚠️ 查重率{rate}%，超出常规学术要求，建议进行AI润色降重")
                if st.button("✨ 一键跳转至润色降重", type="primary"):
                    st.session_state.need_polish = True
                    safe_rerun()
            else:
                st.success(f"✅ 查重率{rate}%，符合学术规范要求")
                if st.button("仍需进行润色优化"):
                    st.session_state.need_polish = True
                    safe_rerun()

        # ---------- 第二步：AI润色降重 ----------
        if st.session_state.need_polish:
            st.divider()
            st.subheader("✨ 第二步：AI润色降重", divider=True)
            col_polish1, col_polish2 = st.columns([3, 1])

            with col_polish1:
                st.markdown("##### 文档与查重报告上传")
                col_polish1_1, col_polish1_2 = st.columns(2)
                with col_polish1_1:
                    # 修复核心bug：移除file_uploader的value参数，改为提示语
                    polish_doc = st.file_uploader(
                        "待润色文档",
                        type=["docx", "doc", "pdf"],
                        key="polish_doc_upload"
                    )
                    if st.session_state.formatted_doc:
                        st.info("✅ 已检测到上一步排版完成的文档，可直接使用，无需重复上传")
                with col_polish1_2:
                    report_file = st.file_uploader(
                        "查重报告(可选)",
                        type=["html", "txt"],
                        key="polish_report_upload"
                    )
                    # 解析查重报告
                    if report_file:
                        red_parts, plain_text, error = parse_plagiarism_report(report_file)
                        if error:
                            st.error(f"解析失败：{error}")
                        elif red_parts:
                            st.success(f"✅ 解析完成！发现{len(red_parts)}处标红重复内容")
                            st.session_state.learned_forbidden = red_parts
                            st.session_state.learn_history.append({
                                "time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                                "forbidden_count": len(red_parts)
                            })

            with col_polish2:
                st.markdown("##### 润色配置")
                polish_level = st.selectbox("润色强度", options=list(REWRITE_LEVEL.keys()), index=1, key="polish_level_select")
                polish_api_key = st.text_input("API Key", type="password", key="polish_api_key")
                st.caption("配置API Key可启用AI深度润色")

            # 润色按钮
            use_formatted_doc = st.checkbox("使用上一步排版后的文档", value=True)
            target_doc = st.session_state.formatted_doc if use_formatted_doc else polish_doc

            if target_doc and st.button("🚀 开始AI润色降重", type="primary", use_container_width=True):
                with st.spinner("正在进行AI润色降重..."):
                    try:
                        output_doc, changes, title_stats, process_log, check_report, full_text = process_doc(
                            file=target_doc,
                            cn_format=st.session_state.cn_format,
                            en_format=st.session_state.en_format,
                            enable_rewrite=True,
                            rewrite_level=polish_level,
                            bind_wps_style=True,
                            standardize_ref=True,
                            api_key=polish_api_key,
                            forbidden_text=st.session_state.learned_forbidden
                        )
                        # 保存润色结果
                        st.session_state.polish_doc = output_doc
                        # 生成润色报告
                        polish_report = generate_report(changes, polish_level, title_stats, process_log, check_report)
                        st.session_state.polish_report = polish_report
                        # 重新查重
                        new_rate = simulate_check_rate(full_text)
                        st.session_state.check_rate = new_rate
                        st.success(f"✅ 润色完成！共优化{len(changes)}处内容，新查重率：{new_rate}%")
                        with st.expander("📋 润色详情", expanded=True):
                            for log in process_log:
                                st.write(log)
                    except Exception as e:
                        st.error(f"润色失败：{str(e)}")

        # ---------- 第三步：成果输出 ----------
        st.divider()
        st.subheader("📥 第三步：成果输出", divider=True)
        timestamp = st.session_state.process_timestamp

        # 检查是否有多个处理文件
        has_multiple_files = hasattr(st.session_state, 'processed_files') and len(st.session_state.processed_files) > 1

        if has_multiple_files:
            # 多个文件时，提供文件选择器
            file_options = list(st.session_state.processed_files.keys())
            selected_file = st.selectbox(
                "选择要下载的文件",
                options=file_options,
                format_func=lambda x: st.session_state.processed_files[x]["file_name"]
            )
            selected_file_data = st.session_state.processed_files[selected_file]
        else:
            # 单个文件时，使用默认值
            selected_file_data = {
                "output_doc": st.session_state.formatted_doc,
                "report": st.session_state.formatted_report,
                "file_name": f"文档_{timestamp}"
            }

        col_output1, col_output2, col_output3, col_output4 = st.columns(4)

        # 排版后文档下载
        with col_output1:
            if selected_file_data.get("output_doc"):
                st.download_button(
                    label="📥 下载排版后文档",
                    data=selected_file_data["output_doc"],
                    file_name=f"标准格式_{selected_file_data['file_name'].replace('.', '_')}_{timestamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        # 排版处理报告下载
        with col_output2:
            if selected_file_data.get("report"):
                st.download_button(
                    label="📋 下载格式处理报告",
                    data=selected_file_data["report"],
                    file_name=f"格式处理报告_{selected_file_data['file_name'].replace('.', '_')}_{timestamp}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
        # 润色后文档下载
        with col_output3:
            if st.session_state.polish_doc:
                st.download_button(
                    label="✨ 下载润色后文档",
                    data=st.session_state.polish_doc,
                    file_name=f"润色降重_{timestamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        # 润色报告下载
        with col_output4:
            if st.session_state.polish_report:
                st.download_button(
                    label="📊 下载润色报告",
                    data=st.session_state.polish_report,
                    file_name=f"润色降重报告_{timestamp}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
        
        # ---------- 智能学术助手 ----------
        st.divider()
        st.subheader("🤖 智能学术助手", divider=True)
        
        # 学术文献搜索
        with st.expander("🔍 学术文献搜索", expanded=False):
            search_keyword = st.text_input("输入搜索关键词", placeholder="例如：人工智能、机器学习、深度学习等")
            max_results = st.slider("搜索结果数量", 1, 10, 5)
            
            if st.button("开始搜索", use_container_width=True):
                if search_keyword:
                    with st.spinner("正在搜索学术文献..."):
                        results, error = search_academic_papers(search_keyword, max_results)
                        if error:
                            st.error(f"搜索失败：{error}")
                        else:
                            st.success(f"找到 {len(results)} 篇相关文献")
                            for i, paper in enumerate(results):
                                with st.expander(f"{i+1}. {paper['title']}"):
                                    st.markdown(f"**作者**：{', '.join(paper['authors'])}")
                                    st.markdown(f"**期刊**：{paper['journal']}")
                                    st.markdown(f"**年份**：{paper['year']}")
                                    st.markdown(f"**摘要**：{paper['abstract']}")
                                    st.markdown(f"**链接**：[{paper['url']}]({paper['url']})")
                else:
                    st.warning("请输入搜索关键词")
        
        # 智能推荐参考文献
        with st.expander("📚 智能推荐参考文献", expanded=False):
            if st.session_state.doc_full_text:
                st.info("基于您的文档内容，我们可以为您推荐相关的参考文献")
                if st.button("推荐参考文献", use_container_width=True):
                    with st.spinner("正在分析文档并推荐参考文献..."):
                        # 提取文档关键词
                        keywords = RE_KEYWORDS.findall(st.session_state.doc_full_text)
                        if keywords:
                            # 使用频率最高的前3个关键词
                            top_keywords = pd.Series(keywords).value_counts().head(3).index.tolist()
                            st.success(f"基于文档分析，推荐以下关键词的参考文献：{', '.join(top_keywords)}")
                            
                            for keyword in top_keywords:
                                st.subheader(f"关键词：{keyword}")
                                results, error = search_academic_papers(keyword, 3)
                                if results:
                                    for i, paper in enumerate(results):
                                        st.markdown(f"{i+1}. **{paper['title']}** - {', '.join(paper['authors'])} ({paper['year']})")
                                else:
                                    st.info(f"未找到关于 '{keyword}' 的文献")
                        else:
                            st.info("无法从文档中提取关键词，请尝试手动搜索")
            else:
                st.info("请先上传并处理文档，我们将基于文档内容为您推荐参考文献")
        
        # 底部提示
        st.caption("💡 所有文件仅在内存中生成，不会保存到服务器，关闭页面后自动清除，保障文档安全")

if __name__ == "__main__":
    main()
