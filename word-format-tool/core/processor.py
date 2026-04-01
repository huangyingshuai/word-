from docx import Document
from io import BytesIO

def is_protected_para(para):
    if not para.text.strip():
        return False
    try:
        if para.part and para.part.type in (1, 2, 3):
            return True
    except:
        pass
    try:
        if para.font.hidden:
            return True
    except:
        pass
    return False

def get_title_level(para, enable_title_regex, last_levels):
    text = para.text.strip()
    if not text:
        return "正文"
    if text.startswith(("一、", "第一章", "1、")):
        return "一级标题"
    elif text.startswith(（"（一）", "1.1", "二、")):
        return "二级标题"
    elif text.startswith(（"1.1.1", "（1）")):
        return "三级标题"
    return "正文"

def process_doc(file, config, number_config, enable_title_regex, force_style, keep_spacing, clear_blank, max_blank):
    doc = Document(BytesIO(file.getvalue()))
    stats = {
        "一级标题": 0, "二级标题": 0, "三级标题": 0,
        "正文": 0, "表格": len(doc.tables),
        "图片": len([r for r in doc.element.xpath('.//a:blip')])
    }
    last_levels = [0, 0, 0]

    def apply_style(para, level):
        style = config[level]
        para.paragraph_format.first_line_indent = style["indent"] * 12700
        if level in ["一级标题", "二级标题", "三级标题"]:
            stats[level] += 1
        else:
            stats["正文"] += 1

    for para in doc.paragraphs:
        if is_protected_para(para):
            continue
        level = get_title_level(para, enable_title_regex, last_levels)
        apply_style(para, level)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if is_protected_para(para):
                        continue
                    apply_style(para, "表格")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output, stats, 1.0, []
